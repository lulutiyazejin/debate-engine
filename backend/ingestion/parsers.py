"""文档解析管道：类型检测 + 六种格式 parser，输出统一 ParsedDocument。

PDF：优先 docling（结构感知），降级 pypdf（纯文本+页码）。
URL：优先 trafilatura，降级 httpx + 正则去标签。
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))


@dataclass
class Section:
    title: str
    text: str
    level: int = 1          # 1=H1/章，2=H2/节
    page_range: str = ""


@dataclass
class ParsedDocument:
    source_type: str        # book/paper/news/url/txt...
    title: str
    author: str | None = None
    year: int | None = None
    sections: list[Section] = field(default_factory=list)
    raw_metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(s.text for s in self.sections)


SKIP_TITLES = ("参考文献", "Bibliography", "References", "索引", "Index",
               "致谢", "Acknowledg", "版权", "目录", "Contents")


def _should_skip(title: str) -> bool:
    return any(k.lower() in title.lower() for k in SKIP_TITLES)


# ---------- 类型检测 ----------
def detect_type(path_or_url: str) -> str:
    s = str(path_or_url).lower()
    if s.startswith(("http://", "https://")):
        return "url"
    ext = Path(s).suffix
    return {".pdf": "pdf", ".docx": "docx", ".doc": "docx",
            ".xlsx": "excel", ".xls": "excel", ".csv": "csv", ".txt": "txt",
            ".md": "md", ".markdown": "md"}.get(ext, "txt")


# ---------- TXT / MD ----------
def parse_txt(path: Path) -> ParsedDocument:
    import chardet
    raw = path.read_bytes()
    enc = chardet.detect(raw[:32768]).get("encoding") or "utf-8"
    text = raw.decode(enc, errors="replace")
    return _text_to_doc(text, title=path.stem, source_type="txt")


def parse_md(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    text = text.replace("\r\n", "\n")  # CRLF 归一化（frontmatter 正则依赖）
    meta: dict = {}
    m = re.match(r"^---\n(.*?)\n---\n", text, re.DOTALL)
    if m:  # frontmatter
        for line in m.group(1).splitlines():
            if ":" in line:
                k, v = line.split(":", 1)
                meta[k.strip()] = v.strip()
        text = text[m.end():]
    doc = _text_to_doc(text, title=meta.get("title", path.stem),
                       source_type=meta.get("source_type", "md"))
    doc.author = meta.get("author")
    doc.year = int(meta["year"]) if str(meta.get("year", "")).isdigit() else None
    doc.raw_metadata = meta
    return doc


def _text_to_doc(text: str, title: str, source_type: str) -> ParsedDocument:
    """按 Markdown 标题切分；无标题则整体一节。"""
    sections: list[Section] = []
    cur_title, cur_level, buf = title, 1, []
    for line in text.splitlines():
        h = re.match(r"^(#{1,3})\s+(.+)$", line)
        if h:
            if buf and "".join(buf).strip():
                sections.append(Section(cur_title, "\n".join(buf).strip(),
                                        cur_level))
            cur_title, cur_level, buf = h.group(2).strip(), len(h.group(1)), []
        else:
            buf.append(line)
    if buf and "".join(buf).strip():
        sections.append(Section(cur_title, "\n".join(buf).strip(), cur_level))
    sections = [s for s in sections if not _should_skip(s.title)]
    if not sections:
        sections = [Section(title, text.strip(), 1)]
    return ParsedDocument(source_type=source_type, title=title,
                          sections=sections)


# ---------- PDF ----------
def parse_pdf(path: Path) -> ParsedDocument:
    try:
        doc = _parse_pdf_docling(path)
    except ImportError:
        doc = _parse_pdf_pypdf(path)
    # 0.1.4 批 6（决策 7）：文字层过薄 → 扫描件 OCR 分支；无组件显式报告
    if len(doc.full_text.strip()) < 200:
        try:
            return _parse_pdf_ocr(path)
        except ImportError:
            doc.raw_metadata["ocr_needed"] = True
            doc.raw_metadata["conversion_note"] = (
                "扫描版 PDF：未安装 OCR 组件（设置→组件中心），"
                "本次只能归档元数据与摘要")
    return doc


def _parse_pdf_ocr(path: Path) -> ParsedDocument:
    """RapidOCR（ONNX，不背框架）+ pypdfium2 渲染；组件中心装入 _extras。"""
    import numpy as np
    import pypdfium2 as pdfium
    from rapidocr_onnxruntime import RapidOCR
    ocr = RapidOCR()
    pdf = pdfium.PdfDocument(str(path))
    pages_text: list[str] = []
    for i in range(len(pdf)):
        bitmap = pdf[i].render(scale=2.0)
        img = np.asarray(bitmap.to_pil())
        result, _ = ocr(img)
        pages_text.append("\n".join(r[1] for r in (result or [])))
    pdf.close()
    sections: list[Section] = []
    step = 8
    for i in range(0, len(pages_text), step):
        body = "\n".join(pages_text[i:i + step]).strip()
        if body:
            sections.append(Section(f"第{i // step + 1}部分", body, 1,
                                    f"p.{i + 1}-{min(i + step, len(pages_text))}"))
    doc = ParsedDocument(source_type="pdf", title=path.stem,
                         sections=sections or [Section(path.stem, "", 1)])
    doc.raw_metadata["parser"] = "rapidocr"
    return doc


def _parse_pdf_docling(path: Path) -> ParsedDocument:
    from docling.document_converter import DocumentConverter
    result = DocumentConverter().convert(str(path))
    md_text = result.document.export_to_markdown()
    doc = _text_to_doc(md_text, title=path.stem, source_type="pdf")
    doc.raw_metadata["parser"] = "docling"
    return doc


def _parse_pdf_pypdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    meta = reader.metadata or {}
    sections: list[Section] = []
    # 优先书签（outline）切章
    try:
        outline = reader.outline
    except Exception:
        outline = []
    bookmarks: list[tuple[str, int]] = []

    def _walk(items):
        for it in items:
            if isinstance(it, list):
                _walk(it)
            else:
                try:
                    bookmarks.append((it.title, reader.get_destination_page_number(it)))
                except Exception:
                    pass
    _walk(outline if isinstance(outline, list) else [])

    pages_text = [(p.extract_text() or "") for p in reader.pages]
    if len(bookmarks) >= 2:
        bookmarks.sort(key=lambda x: x[1])
        for i, (bt, start) in enumerate(bookmarks):
            end = bookmarks[i + 1][1] if i + 1 < len(bookmarks) else len(pages_text)
            body = "\n".join(pages_text[start:end]).strip()
            if body and not _should_skip(bt):
                sections.append(Section(bt, body, 1, f"p.{start+1}-{end}"))
    if not sections:
        # 无书签：每 8 页一节
        step = 8
        for i in range(0, len(pages_text), step):
            body = "\n".join(pages_text[i:i + step]).strip()
            if body:
                sections.append(Section(f"第{i//step+1}部分", body, 1,
                                        f"p.{i+1}-{min(i+step, len(pages_text))}"))
    doc = ParsedDocument(source_type="pdf", title=meta.get("/Title") or path.stem,
                         author=meta.get("/Author"), sections=sections)
    doc.raw_metadata["parser"] = "pypdf"
    return doc


# ---------- DOCX ----------
def parse_docx(path: Path) -> ParsedDocument:
    import docx
    d = docx.Document(str(path))
    sections: list[Section] = []
    cur_title, cur_level, buf = path.stem, 1, []
    for para in d.paragraphs:
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            if buf and "".join(buf).strip():
                sections.append(Section(cur_title, "\n".join(buf).strip(),
                                        cur_level))
            cur_title = para.text.strip() or cur_title
            cur_level = 1 if "1" in style else 2
            buf = []
        else:
            buf.append(para.text)
    if buf and "".join(buf).strip():
        sections.append(Section(cur_title, "\n".join(buf).strip(), cur_level))
    sections = [s for s in sections if not _should_skip(s.title)]
    core = d.core_properties
    return ParsedDocument(source_type="docx", title=core.title or path.stem,
                          author=core.author, sections=sections or
                          [Section(path.stem, "", 1)])


# ---------- EXCEL ----------
def parse_excel(path: Path) -> ParsedDocument:
    """表格 → 文本行描述（AI 转述在入库 Stage 3 由 LLM 处理，此处先结构化）。"""
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sections: list[Section] = []
    for ws in wb.worksheets:
        rows = []
        header: list[str] | None = None
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else str(v) for v in row]
            if not any(vals):
                continue
            if header is None:
                header = vals
                continue
            pairs = [f"{h}={v}" for h, v in zip(header, vals) if v]
            if pairs:
                rows.append("；".join(pairs))
            if len(rows) >= 500:  # 防超大表
                break
        if rows:
            text = f"表格《{ws.title}》共 {len(rows)} 行数据：\n" + "\n".join(rows)
            sections.append(Section(ws.title, text, 1))
    wb.close()
    return ParsedDocument(source_type="excel", title=path.stem,
                          sections=sections or [Section(path.stem, "", 1)])


# ---------- CSV（0.1.4 批 6 决策 8：进支持清单，编码回退） ----------
def parse_csv(path: Path) -> ParsedDocument:
    import csv as _csv
    raw = path.read_bytes()
    text = None
    for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = raw.decode("utf-8", errors="replace")
    rows = list(_csv.reader(text.splitlines()))
    lines = []
    header = rows[0] if rows else []
    for vals in rows[1:501]:
        pairs = [f"{h}={v}" for h, v in zip(header, vals) if v]
        if pairs:
            lines.append("；".join(pairs))
    body = (f"表格《{path.stem}》共 {len(lines)} 行数据：\n" + "\n".join(lines)
            if lines else "")
    return ParsedDocument(source_type="excel", title=path.stem,
                          sections=[Section(path.stem, body, 1)])


# ---------- URL ----------
def parse_url(url: str) -> ParsedDocument:
    html: str | None = None
    try:
        import trafilatura
        html = trafilatura.fetch_url(url)
        text = trafilatura.extract(html) or ""
        meta = trafilatura.extract_metadata(html)
        title = (meta.title if meta else None) or url
        author = meta.author if meta else None
    except ImportError:
        import httpx
        r = httpx.get(url, timeout=20, follow_redirects=True,
                      headers={"User-Agent": "Mozilla/5.0"})
        if r.status_code == 403:
            raise ValueError(f"站点反爬拒绝（403）：{url}；"
                             f"可浏览器下载附件后用文件导入")
        html = r.text
        tm = re.search(r"<title[^>]*>(.*?)</title>", html, re.DOTALL | re.I)
        title = re.sub(r"\s+", " ", tm.group(1)).strip() if tm else url
        author = None
        body = re.sub(r"<(script|style|nav|footer|header)[^>]*>.*?</\1>", " ",
                      html, flags=re.DOTALL | re.I)
        text = re.sub(r"<[^>]+>", " ", body)
        text = re.sub(r"\s{2,}", "\n", text).strip()
    doc = _text_to_doc(text, title=title, source_type="url")
    doc.author = author
    doc.raw_metadata["url"] = url
    _enrich_url_tables(doc, html, url)
    return doc


def _enrich_url_tables(doc: ParsedDocument, html: str | None, url: str) -> None:
    """0.1.4 批 6（决策 8，借 akshare 解析结论）：政府站/统计页表格全抽取
    → md 表格节；附件链接（xls/xlsx/csv/pdf）发现 → raw_metadata 供确认屏提示。
    pandas/lxml 缺位时静默跳过（不阻断正文导入）。"""
    if not html:
        return
    try:
        import io
        import pandas as pd
        tables = pd.read_html(io.StringIO(html))
        for i, t in enumerate(tables):
            if t.shape[0] < 2 or t.shape[1] < 2:
                continue   # 装饰性小表跳过
            md = t.head(200).to_markdown(index=False)
            doc.sections.append(Section(f"页内表格 {i + 1}", md, 2))
    except Exception:
        pass
    try:
        from urllib.parse import urljoin
        links = re.findall(r'href=["\']([^"\']+?\.(?:xlsx?|csv|pdf))["\']',
                           html, flags=re.I)
        if links:
            doc.raw_metadata["attachments"] = sorted(
                {urljoin(url, x) for x in links})[:20]
    except Exception:
        pass


# ---------- 统一入口 ----------
def parse_any(path_or_url: str) -> ParsedDocument:
    t = detect_type(path_or_url)
    if t == "url":
        return parse_url(path_or_url)
    p = Path(path_or_url)
    if not p.exists():
        raise FileNotFoundError(path_or_url)
    return {"pdf": parse_pdf, "docx": parse_docx, "excel": parse_excel,
            "md": parse_md, "txt": parse_txt, "csv": parse_csv}[t](p)
