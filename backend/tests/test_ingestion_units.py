"""嵌入 / 解析 / 分块 测试。"""
from __future__ import annotations

import numpy as np
import pytest

import config
from ingestion.chunker import chunk_document, estimate_tokens
from ingestion.parsers import (ParsedDocument, Section, detect_type,
                               parse_any, parse_md, parse_txt)
from models.embedder import HashEmbedder


class TestEmbedder:
    def test_batch_dim(self):
        emb = HashEmbedder()
        vecs = emb.embed_batch(["市场经济", "计划经济", "自由", "平等", "hello"])
        assert len(vecs) == 5
        assert all(v.shape == (config.EMBEDDING_DIM,) for v in vecs)

    def test_deterministic_and_similarity(self):
        emb = HashEmbedder()
        a1 = emb.embed("市场经济依靠价格信号")
        a2 = emb.embed("市场经济依靠价格信号")
        b = emb.embed("今天天气很好适合郊游")
        assert np.allclose(a1, a2)
        sim_same = float(a1 @ a2)
        sim_diff = float(a1 @ b)
        assert sim_same > sim_diff  # 同文本相似度必须更高


class TestParsers:
    def test_detect_type(self):
        assert detect_type("a.pdf") == "pdf"
        assert detect_type("b.docx") == "docx"
        assert detect_type("https://example.com/x") == "url"
        assert detect_type("c.md") == "md"
        assert detect_type("noext") == "txt"

    def test_parse_txt_gbk(self, tmp_path):
        p = tmp_path / "gbk.txt"
        p.write_bytes("市场经济与价格信号的关系研究".encode("gbk"))
        doc = parse_txt(p)
        assert "价格信号" in doc.full_text

    def test_parse_md_frontmatter_and_headings(self, tmp_path):
        p = tmp_path / "t.md"
        p.write_text("---\ntitle: 试论自由\nauthor: 李四\nyear: 1999\n---\n"
                     "# 第一章\n正文A\n## 小节\n正文B\n# 参考文献\n[1] xxx\n",
                     encoding="utf-8")
        doc = parse_md(p)
        assert doc.title == "试论自由"
        assert doc.author == "李四" and doc.year == 1999
        titles = [s.title for s in doc.sections]
        assert "第一章" in titles and "参考文献" not in titles  # 过滤规则

    def test_parse_docx(self, tmp_path):
        import docx
        d = docx.Document()
        d.add_heading("引言", level=1)
        d.add_paragraph("这是引言内容。")
        d.add_heading("方法", level=2)
        d.add_paragraph("这是方法内容。")
        p = tmp_path / "t.docx"
        d.save(str(p))
        doc = parse_any(str(p))
        assert doc.source_type == "docx"
        assert "引言内容" in doc.full_text

    def test_parse_excel(self, tmp_path):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["国家", "GDP"])
        ws.append(["A国", 100])
        p = tmp_path / "t.xlsx"
        wb.save(str(p))
        doc = parse_any(str(p))
        assert "GDP=100" in doc.full_text

    def test_missing_file_raises(self):
        with pytest.raises(FileNotFoundError):
            parse_any("Z:/not/exist/file.txt")


class TestChunker:
    def test_short_doc_single_chunk(self):
        doc = ParsedDocument(source_type="txt", title="短文",
                             sections=[Section("短文", "很短的内容。", 1)])
        chunks = chunk_document(doc)
        assert len(chunks) == 1
        assert chunks[0].chapter_title == "短文"

    def test_chapter_boundary(self):
        secs = [Section(f"第{i}章", "内容" * 800, 1) for i in range(1, 4)]
        doc = ParsedDocument(source_type="pdf", title="书", sections=secs)
        chunks = chunk_document(doc, short_doc_tokens=100)
        assert len(chunks) == 3
        assert chunks[0].chapter_title == "第1章"

    def test_oversize_chapter_split(self):
        long_text = "\n\n".join("这是一个足够长的段落。" * 50 for _ in range(20))
        doc = ParsedDocument(source_type="pdf", title="书",
                             sections=[Section("超长章", long_text, 1)])
        chunks = chunk_document(doc, max_tokens=2000, short_doc_tokens=100)
        assert len(chunks) > 1
        assert all(c.token_count <= 2000 * 1.1 for c in chunks)

    def test_estimate_tokens(self):
        assert estimate_tokens("中文十个字中文十个字") == 10
        assert estimate_tokens("") == 0
